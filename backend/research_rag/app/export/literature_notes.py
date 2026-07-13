from __future__ import annotations
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from docx import Document

@dataclass
class CitationItem:
    label: str
    quote: str
    page_start: Optional[int] = None
    page_end: Optional[int] = None

@dataclass
class LiteratureNote:
    title: str
    question: str
    answer: str
    key_findings: List[str] = field(default_factory=list)
    compared_papers: List[str] = field(default_factory=list)
    citations: List[CitationItem] = field(default_factory=list)
    limitations: List[str] = field(default_factory=list)
    created_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())

class LiteratureNotesExporter:
    @staticmethod
    def to_markdown(note: LiteratureNote) -> str:
        lines = [
            f"# {note.title}",
            "",
            f"**Question:** {note.question}",
            "",
            "## Answer",
            note.answer,
            "",
        ]

        if note.key_findings:
            lines.extend(["## Key Findings", ""])
            for item in note.key_findings:
                lines.append(f"- {item}")
            lines.append("")

        if note.compared_papers:
            lines.extend(["## Compared Papers", ""])
            for paper in note.compared_papers:
                lines.append(f"- {paper}")
            lines.append("")

        if note.citations:
            lines.extend(["## Citations", ""])
            for i, c in enumerate(note.citations, start=1):
                pages = ""
                if c.page_start is not None:
                    pages = f" (p. {c.page_start}" + (f"-{c.page_end}" if c.page_end and c.page_end != c.page_start else "") + ")"
                lines.append(f"{i}. **{c.label}**{pages}")
                lines.append(f"   - {c.quote}")
            lines.append("")

        if note.limitations:
            lines.extend(["## Limitations / Uncertainty", ""])
            for item in note.limitations:
                lines.append(f"- {item}")
            lines.append("")

        lines.extend(["---", f"_Generated on: {note.created_at}_", ""])
        return "\n".join(lines)

    @staticmethod
    def save_markdown(note: LiteratureNote, output_path: str | Path) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(LiteratureNotesExporter.to_markdown(note), encoding="utf-8")
        return output_path

    @staticmethod
    def save_docx(note: LiteratureNote, output_path: str | Path) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(note.title, level=0)

        doc.add_heading("Question", level=1)
        doc.add_paragraph(note.question)

        doc.add_heading("Answer", level=1)
        doc.add_paragraph(note.answer)

        if note.key_findings:
            doc.add_heading("Key Findings", level=1)
            for item in note.key_findings:
                doc.add_paragraph(item, style="List Bullet")

        if note.compared_papers:
            doc.add_heading("Compared Papers", level=1)
            for item in note.compared_papers:
                doc.add_paragraph(item, style="List Bullet")

        if note.citations:
            doc.add_heading("Citations", level=1)
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Source"
            hdr[1].text = "Pages"
            hdr[2].text = "Supporting Quote"

            for c in note.citations:
                row = table.add_row().cells
                row[0].text = c.label
                if c.page_start is not None:
                    if c.page_end and c.page_end != c.page_start:
                        row[1].text = f"{c.page_start}-{c.page_end}"
                    else:
                        row[1].text = str(c.page_start)
                else:
                    row[1].text = ""
                row[2].text = c.quote

        if note.limitations:
            doc.add_heading("Limitations / Uncertainty", level=1)
            for item in note.limitations:
                doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph(f"Generated on: {note.created_at}")
        doc.save(str(output_path))
        return output_path
